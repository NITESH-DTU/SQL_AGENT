import os
import pandas as pd
import pdfplumber
import docx
from typing import Dict, Any, List, Optional

class FileImporter:
    def __init__(self, db_manager=None):
        self.db_manager = db_manager

    def get_preview(self, file_path: str) -> Dict[str, Any]:
        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext == '.csv':
                df = pd.read_csv(file_path, nrows=5)
                return {"type": "table", "data": df.to_dict(orient='records'), "columns": df.columns.tolist()}
            elif ext in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path, nrows=5)
                return {"type": "table", "data": df.to_dict(orient='records'), "columns": df.columns.tolist()}
            elif ext == '.pdf':
                with pdfplumber.open(file_path) as pdf:
                    first_page = pdf.pages[0]
                    text = first_page.extract_text()
                    tables = first_page.extract_tables()
                    if tables:
                        df = pd.DataFrame(tables[0][1:], columns=tables[0][0])
                        return {"type": "table", "data": df.head(5).to_dict(orient='records'), "columns": df.columns.tolist()}
                    return {"type": "text", "preview": text[:1000] if text else "No text found"}
            elif ext == '.docx':
                doc = docx.Document(file_path)
                if doc.tables:
                    table = doc.tables[0]
                    keys = [cell.text for cell in table.rows[0].cells]
                    data = []
                    for row in table.rows[1:6]:
                        data.append([cell.text for cell in row.cells])
                    df = pd.DataFrame(data, columns=keys)
                    return {"type": "table", "data": df.to_dict(orient='records'), "columns": df.columns.tolist()}
                text = "\n".join([p.text for p in doc.paragraphs[:20]])
                return {"type": "text", "preview": text[:1000]}
            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    lines = f.readlines()
                    preview = "".join(lines[:20])
                    return {"type": "text", "preview": preview[:1000]}
            return {"error": "Unsupported file type"}
        except Exception as e:
            return {"error": str(e)}

    def import_to_table(self, file_path: str, table_name: str, mode: str) -> Dict[str, Any]:
        if not self.db_manager:
            return {"error": "DB Manager not initialized"}
            
        ext = os.path.splitext(file_path)[1].lower()
        df = None
        try:
            if ext == '.csv':
                df = pd.read_csv(file_path)
            elif ext in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path)
            elif ext == '.pdf':
                all_tables = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_tables = page.extract_tables()
                        for t in page_tables:
                            if t and len(t) > 1:
                                all_tables.append(pd.DataFrame(t[1:], columns=t[0]))
                if all_tables:
                    df = pd.concat(all_tables, ignore_index=True)
                else:
                    with pdfplumber.open(file_path) as pdf:
                        text = "\n".join([p.extract_text() or "" for p in pdf.pages])
                        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                        df = pd.DataFrame(paragraphs, columns=['content'])
            elif ext == '.docx':
                doc = docx.Document(file_path)
                all_tables = []
                for table in doc.tables:
                    if len(table.rows) > 1:
                        keys = [cell.text for cell in table.rows[0].cells]
                        data = [[cell.text for cell in row.cells] for row in table.rows[1:]]
                        all_tables.append(pd.DataFrame(data, columns=keys))
                if all_tables:
                    df = pd.concat(all_tables, ignore_index=True)
                else:
                    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
                    df = pd.DataFrame(paragraphs, columns=['content'])
            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    lines = [line.strip() for line in f.readlines() if line.strip()]
                    df = pd.DataFrame(lines, columns=['content'])
            
            if df is not None:
                # Sanitize column names: lowercase, underscore, alphanumeric only
                new_cols = []
                for c in df.columns:
                    sc = "".join([char if char.isalnum() else "_" for char in str(c).lower()])
                    while "__" in sc: sc = sc.replace("__", "_")
                    new_cols.append(sc.strip("_"))
                df.columns = new_cols
                
                # Replace NaN with None for SQL compatibility
                df = df.where(pd.notnull(df), None)
                
                if mode == "create_new":
                    cols_sql = []
                    for col, dtype in df.dtypes.items():
                        sql_type = "TEXT"
                        if "int" in str(dtype): sql_type = "INTEGER"
                        elif "float" in str(dtype): sql_type = "REAL"
                        cols_sql.append(f'"{col}" {sql_type}')
                    
                    create_sql = f"CREATE TABLE \"{table_name}\" ({', '.join(cols_sql)})"
                    self.db_manager.execute_write(create_sql)
                
                # Bulk insert
                placeholders = ", ".join(["?" if self.db_manager.db_type == "sqlite" else "%s"] * len(df.columns))
                cols_str = ", ".join([f'"{c}"' for c in df.columns])
                insert_sql = f"INSERT INTO \"{table_name}\" ({cols_str}) VALUES ({placeholders})"
                
                rows_to_insert = [tuple(x) for x in df.values]
                chunk_size = 500
                total_inserted = 0
                for i in range(0, len(rows_to_insert), chunk_size):
                    batch = rows_to_insert[i:i+chunk_size]
                    res = self.db_manager.execute_many(insert_sql, batch)
                    if "error" in res:
                        return {"error": f"Import failed at batch {i}: {res['error']}"}
                    total_inserted += len(batch)
                
                return {"status": "success", "rows_imported": total_inserted, "table": table_name}
            
            return {"error": "No data found to import"}
        except Exception as e:
            return {"error": str(e)}
